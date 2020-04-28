#--read and write to a .docx file 
#---------
Python can create and modify Word documents, which have 
the .docx file extension, with the python-docx module. 

You can install the module by running 
pip install python-docx
#---------
this will install 2 packages : lxml and python-docx
#---------
#--use file "demo.docx" for this code

import docx
fileName = 'suven\demo.docx'
doc = docx.Document(fileName)
len_of_docx = len(doc.paragraphs)
print(f"#paras in {fileName} is {len_of_docx}")

#--The Document object contains a list of Paragraph objects for the paragraphs in the document.

#--python-docx lib treats the entire word doc as DOM
#where Document is the root and refers to entire doc
#and every element there-after is a child to the root
line_1 = doc.paragraphs[0].text
print(f"first Line of {fileName} is {line_1}")

#then the second line would be 
line_2 = doc.paragraphs[1].text
print(f"Next Line of {fileName} is {line_2}")

#--Each of these Paragraph objects contains a list of one or more Run objects. A Run object is a contiguous run of text with the same style.
num_of_textStyles_in_second_line = len(doc.paragraphs[1].runs)
print(f"number of runs in second line is {num_of_textStyles_in_second_line}")


#--printing all runs in 2nd para
for i in range(0,num_of_textStyles_in_second_line):
 print(doc.paragraphs[1].runs[i].text)

print("--------------------------------------")
print("reading all text contents in the file")
#--defining a fn to read the entire file as text
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(' ' + para.text)
    return '\n'.join(fullText)

#--call the above function
print(getText(fileName))


#-----------------
#--create a new doc, see styles  
import docx

#--create a new doc and write some thing to it
print("--------------------")
print("make a new doc file from code")
doc = docx.Document()
doc.add_paragraph('Hello world!', 'Title')
doc.add_paragraph('Word doc scrapping!')
#add_heading('String data', int Hdg) 
#int Hdg can be 0,1,2,3,4. 
doc.add_heading('Dummy Test File', 0)
doc.save('suven\helloworld.docx')


#--go to word doc press CTRL-ALT-SHIFT-S
#--would see all styles 
#--we can read styles or change them from our code
New_fileName = 'suven\helloworld.docx'
doc = docx.Document(New_fileName)
print(doc.paragraphs[1].text)

#--style refers to style id
print(doc.paragraphs[1].style) 


#--------------------------------
#--change the style of a para from above doc--
import docx

New_fileName = 'suven\helloworld.docx'
doc = docx.Document(New_fileName)

#--lets change the style. use style_name   
doc.paragraphs[1].style.name = 'Quote'
#try different styles : 'Normal' , 'Heading5', 'ListBullet', 'Quote', 'Title'

doc.save(New_fileName)
print(f"see {New_fileName} for change in style")

#----------------------------------

#----------------------------------
Note to extract text, hyperlinks, images* use docxpy
#----------------------------------
pip install docxpy
#----------------------------------
#--below code reads text, links, images* from docx
 
import docxpy

file = 'suven\One minute read on Web scraping.docx'

# extract text
text = docxpy.process(file)
print(text)
print("-------------------------------")

# extract hyperlinks
doc = docxpy.DOCReader(file)
doc.process()  # process file
hyperlinks = doc.data['links']
print(hyperlinks)


#the same package docxpy can be used for extracting images. But code fails most of the times.
# extract text and write images to /suven/img_dir
#text = docxpy.process(file, "suven\images") 
#print("only text part : ", text)

#------------------------------
#-- hence for extracting images use docx2txt
#---------------
pip install docx2txt 
#---------------

import docx2txt
file = 'suven\One minute read on Web scraping.docx'
docx2txt.process(file,  "C:\Program Files\Python36\suven\images")

#-------------------------------
#-------------------------------

#--Reading from PDF doc
#--------------------
pip install PyPDF2
#--------------------
Note : PyPDF2 does not have a way to extract images, charts, or other media from PDF documents, but it can extract text and return it as a Python string.
#--------------------

import PyPDF2

pdfFileObj = open('suven\sample.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
print("Number of pages : ", pdfReader.numPages)

#reads first page contents. 
#Content quite close to actual. 
pageObj = pdfReader.getPage(0)
print(pageObj.extractText())

#try exceeding the page_index_no. beyond PDF len
#pageObj = pdfReader.getPage(1)

#-------------------------
#-- creating a new combined pdf from 2 pdf docs
import PyPDF2

#be careful don't use \. use / in the path
pdf1File = open('suven/pointers on Resume.pdf', 'rb')
pdf2File = open('suven/resume.pdf', 'rb')

pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
pdf2Reader = PyPDF2.PdfFileReader(pdf2File)

pdfWriter = PyPDF2.PdfFileWriter()

for pageNum in range(pdf1Reader.numPages):
 pageObj = pdf1Reader.getPage(pageNum)
 pdfWriter.addPage(pageObj)

for pageNum in range(pdf2Reader.numPages):
 pageObj = pdf2Reader.getPage(pageNum)
 pdfWriter.addPage(pageObj)

pdfOutputFile = open('suven/combined.pdf', 'wb')
pdfWriter.write(pdfOutputFile)
pdfOutputFile.close()
pdf1File.close()
pdf2File.close()

#------------------------------
Note : PyPDF2 cannot insert pages in the middle of a PdfFileWriter object; the addPage() method will only add pages to the end.
#------------------------------

#----Rotating Pages of a pdf document-----
import PyPDF2

#--be careful use / , not \
orginal = open('suven/resume.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(orginal)
page = pdfReader.getPage(0)
page.rotateClockwise(90)

pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(page)
resultPdfFile = open('suven/rotatedResume.pdf', 'wb')
pdfWriter.write(resultPdfFile)
resultPdfFile.close()
orginal.close()

#--------------------------------
#--------------------------------

#--Add watermark to a pdf doc-------
import PyPDF2
original = open('suven/rotatedResume.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(original)

originalFirstPage = pdfReader.getPage(0)

pdfWatermarkReader = PyPDF2.PdfFileReader(open('suven/watermark.pdf', 'rb'))
originalFirstPage.mergePage(pdfWatermarkReader.getPage(0))

pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(originalFirstPage)

resultPdfFile = open('suven/watermarked_Resume.pdf', 'wb')
pdfWriter.write(resultPdfFile)
original.close()
resultPdfFile.close()

#-------------------------

#----Encrypting PDFs, i.e password protecting--
import PyPDF2
pdfFile = open('suven/resume.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFile)
pdfWriter = PyPDF2.PdfFileWriter()

for pageNum in range(pdfReader.numPages):
 pdfWriter.addPage(pdfReader.getPage(pageNum))

#encrypt means making password protected 
pdfWriter.encrypt('SuVeN')

resultPdf = open('suven/encryptedResume.pdf', 'wb')
pdfWriter.write(resultPdf)
resultPdf.close()

#-----------------------

#--decrypting a pdf doc and reading from it
import PyPDF2

pdfReader = PyPDF2.PdfFileReader(open('suven/encryptedResume.pdf', 'rb'))

if pdfReader.isEncrypted :
 pdfReader.decrypt('SuVeN') #try wrong password
 
pageObj = pdfReader.getPage(0)
print(pageObj.extractText())

#------------------------
